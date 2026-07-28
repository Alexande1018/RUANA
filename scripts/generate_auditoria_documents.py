#!/usr/bin/env python3
"""Generate PDF and DOCX exports from the RUANA forensic audit markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "AUDITORIA_FORENSE_RUANA.md"
OUT_DIR = ROOT / "docs" / "exports"
FONT_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
FONT_MONO = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("🟢", "[ESENCIAL]").replace("🟡", "[DUDOSO]").replace("🔴", "[ELIMINABLE]")


def parse_blocks(lines: list[str]) -> list[dict]:
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            blocks.append({"type": "hr"})
            i += 1
            continue
        if line.startswith("```"):
            lang = line.strip("` ").strip() or "text"
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines)})
            i += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                if re.match(r"^\|\s*[-:]+", tl):
                    continue
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            blocks.append({"type": "table", "rows": rows})
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            blocks.append({"type": "heading", "level": level, "text": line[level:].strip()})
            i += 1
            continue
        if line.startswith("- "):
            items: list[str] = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(lines[i][2:].strip())
                i += 1
            blocks.append({"type": "list", "items": items})
            continue
        if re.match(r"^\d+\.\s", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i]):
                items.append(re.sub(r"^\d+\.\s*", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "olist", "items": items})
            continue
        if not line.strip():
            i += 1
            continue
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "- ", "|", "```")) and not re.match(r"^\d+\.\s", lines[i]) and lines[i].strip() != "---":
            para_lines.append(lines[i])
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines).strip()})
    return blocks


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            clean = strip_md_inline(part)
            p.add_run(clean)


def build_docx(blocks: list[dict], output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            level = min(block["level"], 4)
            doc.add_heading(strip_md_inline(block["text"]), level=level)
        elif btype == "paragraph":
            add_rich_paragraph(doc, block["text"])
        elif btype == "list":
            for item in block["items"]:
                add_rich_paragraph(doc, item, style="List Bullet")
        elif btype == "olist":
            for item in block["items"]:
                add_rich_paragraph(doc, item, style="List Number")
        elif btype == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(8)
        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell_text = strip_md_inline(row[c_idx]) if c_idx < len(row) else ""
                    table.rows[r_idx].cells[c_idx].text = cell_text
            doc.add_paragraph("")
        elif btype == "hr":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.LINE)
    doc.save(output)


class AuditPDF(FPDF):
    def __init__(self) -> None:
        super().__init__(format="A4", unit="mm")
        self.set_auto_page_break(auto=True, margin=15)
        self.add_font("DejaVu", "", FONT_REGULAR)
        self.add_font("DejaVu", "B", FONT_BOLD)
        self.add_font("DejaVuMono", "", FONT_MONO)

    def write_heading(self, text: str, level: int) -> None:
        sizes = {1: 18, 2: 14, 3: 12, 4: 11}
        self.ln(2 if level > 1 else 4)
        self.set_font("DejaVu", "B", sizes.get(level, 10))
        self.multi_cell(0, 7, strip_md_inline(text))
        self.ln(1)

    def write_paragraph(self, text: str, mono: bool = False) -> None:
        self.set_font("DejaVuMono" if mono else "DejaVu", "", 8 if mono else 9)
        self.multi_cell(0, 4.5, strip_md_inline(text))
        self.ln(1)

    def write_table(self, rows: list[list[str]]) -> None:
        if not rows:
            return
        cols = max(len(r) for r in rows)
        widths = [190 / cols] * cols
        self.set_font("DejaVu", "", 7)
        for r_idx, row in enumerate(rows):
            if r_idx == 0:
                self.set_font("DejaVu", "B", 7)
            else:
                self.set_font("DejaVu", "", 7)
            x0 = self.get_x()
            y0 = self.get_y()
            max_h = 0
            cell_lines: list[list[str]] = []
            for c_idx in range(cols):
                text = strip_md_inline(row[c_idx]) if c_idx < len(row) else ""
                lines = self.multi_cell(widths[c_idx], 4, text, split_only=True)
                cell_lines.append(lines)
                max_h = max(max_h, len(lines) * 4)
            if y0 + max_h > 280:
                self.add_page()
                y0 = self.get_y()
            for c_idx in range(cols):
                x = x0 + sum(widths[:c_idx])
                self.set_xy(x, y0)
                self.multi_cell(widths[c_idx], 4, "\n".join(cell_lines[c_idx]), border=1)
            self.set_xy(x0, y0 + max_h)
        self.ln(2)


def build_pdf(blocks: list[dict], output: Path) -> None:
    pdf = AuditPDF()
    pdf.add_page()
    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            pdf.write_heading(block["text"], block["level"])
        elif btype == "paragraph":
            pdf.write_paragraph(block["text"])
        elif btype in {"list", "olist"}:
            for idx, item in enumerate(block["items"], 1):
                prefix = "- " if btype == "list" else f"{idx}. "
                pdf.write_paragraph(prefix + item)
        elif btype == "code":
            pdf.write_paragraph(block["text"], mono=True)
        elif btype == "table":
            pdf.write_table(block["rows"])
        elif btype == "hr":
            pdf.ln(2)
    pdf.output(str(output))


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    text = SOURCE.read_text(encoding="utf-8")
    blocks = parse_blocks(text.splitlines())
    docx_path = OUT_DIR / "AUDITORIA_FORENSE_RUANA.docx"
    pdf_path = OUT_DIR / "AUDITORIA_FORENSE_RUANA.pdf"
    build_docx(blocks, docx_path)
    build_pdf(blocks, pdf_path)
    print(f"DOCX: {docx_path}")
    print(f"PDF:  {pdf_path}")
    print(f"DOCX size: {docx_path.stat().st_size:,} bytes")
    print(f"PDF size:  {pdf_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
